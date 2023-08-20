# ---
# jupyter:
#   jupytext:
#     formats: ipynb,py:light
#     text_representation:
#       extension: .py
#       format_name: light
#       format_version: '1.5'
#       jupytext_version: 1.14.6
#   kernelspec:
#     display_name: Python 3 (ipykernel)
#     language: python
#     name: python3
# ---

# # Ingest website to FAISS

# ## Install/ import stuff we need

import os
from pathlib import Path
import re
import requests
import pandas as pd
import dateutil.parser
from typing import TypeVar, List

from langchain.embeddings import HuggingFaceInstructEmbeddings, HuggingFaceEmbeddings
from langchain.vectorstores.faiss import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from langchain.document_loaders import PyPDFLoader

import magic
from bs4 import BeautifulSoup
from docx import Document as Doc
from pypdf import PdfReader
from docx import Document

PandasDataFrame = TypeVar('pd.core.frame.DataFrame')
# -

split_strat = [".", "!", "?", "\n\n", "\n", ",", " ", ""]
chunk_size = 1000
chunk_overlap = 200

## Overarching ingest function:


def determine_file_type(file_path):
    """
    Determine the MIME type of the given file using the magic library.
    
    Parameters:
        file_path (str): Path to the file.
    
    Returns:
        str: MIME type of the file.
    """
    return magic.from_file(file_path, mime=True)

def parse_pdf(file) -> List[str]:

    """
    Extract text from a PDF file.
    
    Parameters:
        file_path (str): Path to the PDF file.
    
    Returns:
        List[str]: Extracted text from the PDF.
    """
    
    output = []
    for i in range(0,len(file)):
        print(file[i].name)
        pdf = PdfReader(file[i].name) #[i]
        for page in pdf.pages:
            text = page.extract_text()
            # Merge hyphenated words
            text = re.sub(r"(\w+)-\n(\w+)", r"\1\2", text)
            # Fix newlines in the middle of sentences
            text = re.sub(r"(?<!\n\s)\n(?!\s\n)", " ", text.strip())
            # Remove multiple newlines
            text = re.sub(r"\n\s*\n", "\n\n", text)
            output.append(text)
    return output


def parse_docx(file_path):
    """
    Reads the content of a .docx file and returns it as a string.

    Parameters:
    - file_path (str): Path to the .docx file.

    Returns:
    - str: Content of the .docx file.
    """
    doc = Doc(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


def parse_txt(file_path):
    """
    Read text from a TXT or HTML file.
    
    Parameters:
        file_path (str): Path to the TXT or HTML file.
    
    Returns:
        str: Text content of the file.
    """
    with open(file_path, 'r', encoding="utf-8") as file:
        return file.read()



def parse_file(file_paths):
    """
    Accepts a list of file paths, determines each file's type, 
    and passes it to the relevant parsing function.
    
    Parameters:
        file_paths (list): List of file paths.
    
    Returns:
        dict: A dictionary with file paths as keys and their parsed content (or error message) as values.
    """
    if not isinstance(file_paths, list):
        raise ValueError("Expected a list of file paths.")
    
    mime_type_to_parser = {
        'application/pdf': parse_pdf,
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': parse_docx,
        'text/plain': parse_txt,
        'text/html': parse_html
    }
    
    parsed_contents = {}

    for file_path in file_paths:
        mime_type = determine_file_type(file_path)
        if mime_type in mime_type_to_parser:
            parsed_contents[file_path] = mime_type_to_parser[mime_type](file_path)
        else:
            parsed_contents[file_path] = f"Unsupported file type: {mime_type}"
    
    return parsed_contents




def parse_html(page_url, div_filter="p"):
    """
    Determine if the source is a web URL or a local HTML file, extract the content based on the div of choice. Also tries to extract dates (WIP)

    Parameters:
        page_url (str): The web URL or local file path.

    Returns:
        str: Extracted content.
    """

    def is_web_url(s):
    """
    Check if the input string is a web URL.
    """
    return s.startswith("http://") or s.startswith("https://")

    def is_local_html_file(s):
    """
    Check if the input string is a path to a local HTML file.
    """
    return (s.endswith(".html") or s.endswith(".htm")) and os.path.isfile(s)

    def extract_text_from_source(source):
    """
    Determine if the source is a web URL or a local HTML file, 
    and then extract its content accordingly.

    Parameters:
        source (str): The web URL or local file path.

    Returns:
        str: Extracted content.
    """
        if is_web_url(source):
            response = requests.get(source)
            response.raise_for_status()  # Raise an HTTPError for bad responses
            return response.text
        elif is_local_html_file(source):
            with open(source, 'r', encoding='utf-8') as file:
                return file.read()
        else:
            raise ValueError("Input is neither a valid web URL nor a local HTML file path.")

    def clean_html_data(data, date_filter="", div_filt="p"):
    """
    Extracts and cleans data from HTML content.

    Parameters:
        data (str): HTML content to be parsed.
        date_filter (str, optional): Date string to filter results. If set, only content with a date greater than this will be returned.
        div_filt (str, optional): HTML tag to search for text content. Defaults to "p".

    Returns:
        tuple: Contains extracted text and date as strings. Returns empty strings if not found.
    """
    
        soup = BeautifulSoup(data, 'html.parser')

        # Function to exclude div with id "bar"
        def exclude_div_with_id_bar(tag):
            return tag.has_attr('id') and tag['id'] == 'related-links'

        text_elements = soup.find_all(div_filt)
        date_elements = soup.find_all(div_filt, {"class": "page-neutral-intro__meta"})
    
        # Extract date
        date_out = ""
        if date_elements:
            date_out = re.search(">(.*?)<", str(date_elements[0])).group(1)
            date_dt = dateutil.parser.parse(date_out)

            if date_filter:
                date_filter_dt = dateutil.parser.parse(date_filter)
                if date_dt < date_filter_dt:
                    return '', date_out

        # Extract text
        text_out_final = ""
        if text_elements:
            text_out_final = '\n'.join(paragraph.text for paragraph in text_elements)
        else:
            print(f"No elements found with tag '{div_filt}'. No text returned.")
    
        return text_out_final, date_out
    

    #page_url = "https://pypi.org/project/InstructorEmbedding/" #'https://www.ons.gov.uk/visualisations/censusareachanges/E09000022/index.html'

    html_text = extract_text_from_source(page_url)
    #print(page.text)

    texts = []
    metadatas = []

    clean_text, date = clean_html_data(html_text, date_filter="", div_filt=div_filter)
    texts.append(clean_text)
    metadatas.append({"source": page_url, "date":str(date)})

    return texts, metadatas


# +
# Convert parsed text to docs
# -

def text_to_docs(text_dict: dict, chunk_size: int = chunk_size) -> List[Document]:
    """
    Converts the output of parse_file (a dictionary of file paths to content)
    to a list of Documents with metadata.
    """
    
    doc_chunks = []

    for file_path, content in text_dict.items():
        ext = os.path.splitext(file_path)[1].lower()

        # Depending on the file extension, handle the content
        if ext == '.pdf':
            docs = pdf_text_to_docs(content, chunk_size)
        elif ext in ['.html', '.htm', '.txt', '.docx']:
            # Assuming you want to process HTML similarly to PDF in this context
            docs = html_text_to_docs(content, chunk_size)
        else:
            print(f"Unsupported file type {ext} for {file_path}. Skipping.")
            continue

        # Add filename as metadata
        for doc in docs:
            doc.metadata["file"] = file_path
        
        doc_chunks.extend(docs)

    return doc_chunks



def pdf_text_to_docs(text: str, chunk_size: int = chunk_size) -> List[Document]:
    """Converts a string or list of strings to a list of Documents
    with metadata."""
    if isinstance(text, str):
        # Take a single string as one page
        text = [text]
        
    page_docs = [Document(page_content=page) for page in text]

    # Add page numbers as metadata
    for i, doc in enumerate(page_docs):
        doc.metadata["page"] = i + 1

    # Split pages into chunks
    doc_chunks = []

    for doc in page_docs:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            separators=split_strat,#["\n\n", "\n", ".", "!", "?", ",", " ", ""],
            chunk_overlap=chunk_overlap,
        )
        chunks = text_splitter.split_text(doc.page_content)

        
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk, metadata={"page": doc.metadata["page"], "chunk": i}
            )
            # Add sources a metadata
            doc.metadata["page_chunk"] = f"{doc.metadata['page']}-{doc.metadata['chunk']}"
            doc_chunks.append(doc)
    return doc_chunks

def html_text_to_docs(texts, metadatas, chunk_size:int = chunk_size):

    text_splitter = RecursiveCharacterTextSplitter(
        separators=split_strat,#["\n\n", "\n", ".", "!", "?", ",", " ", ""],
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len
    )

    #print(texts)
    #print(metadatas)

    documents = text_splitter.create_documents(texts, metadatas=metadatas)

    for i, chunk in enumerate(documents):
        chunk.metadata["chunk"] = i + 1

    return documents




    

# # Functions for working with documents after loading them back in

def pull_out_data(series):

    # define a lambda function to convert each string into a tuple
    to_tuple = lambda x: eval(x)

    # apply the lambda function to each element of the series
    series_tup = series.apply(to_tuple)

    series_tup_content = list(zip(*series_tup))[1]

    series = pd.Series(list(series_tup_content))#.str.replace("^Main post content", "", regex=True).str.strip()

    return series


def docs_from_csv(df):

    import ast
    
    documents = []
    
    page_content = pull_out_data(df["0"])
    metadatas = pull_out_data(df["1"])

    for x in range(0,len(df)):       
        new_doc = Document(page_content=page_content[x], metadata=metadatas[x])
        documents.append(new_doc)
        
    return documents


def docs_from_lists(docs, metadatas):

    documents = []

    for x, doc in enumerate(docs):
        new_doc = Document(page_content=doc, metadata=metadatas[x])
        documents.append(new_doc)
        
    return documents


def docs_elements_from_csv_save(docs_path="documents.csv"):

    documents = pd.read_csv(docs_path)

    docs_out = docs_from_csv(documents)

    out_df = pd.DataFrame(docs_out)

    docs_content = pull_out_data(out_df[0].astype(str))

    docs_meta = pull_out_data(out_df[1].astype(str))

    doc_sources = [d['source'] for d in docs_meta]

    return out_df, docs_content, docs_meta, doc_sources


# documents = html_text_to_docs(texts, metadatas)
#
# documents[0]
#
# pd.DataFrame(documents).to_csv("documents.csv", index=None)

# ## Create embeddings and save faiss vector store to the path specified in `save_to`

def load_embeddings(model_name = "hkunlp/instructor-large"):

    if model_name == "hkunlp/instructor-large":
        embeddings_func = HuggingFaceInstructEmbeddings(model_name=model_name,
        embed_instruction="Represent the paragraph for retrieval: ",
        query_instruction="Represent the question for retrieving supporting documents: "
        )

    else: 
        embeddings_func = HuggingFaceEmbeddings(model_name=model_name)

    global embeddings

    embeddings = embeddings_func

    #return embeddings_func


def embed_faiss_save_to_zip(docs_out, save_to="faiss_lambeth_census_embedding", model_name = "hkunlp/instructor-large"):

    load_embeddings(model_name=model_name)

    #embeddings_fast = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

    print(f"> Total split documents: {len(docs_out)}")

    vectorstore = FAISS.from_documents(documents=docs_out, embedding=embeddings)
        

    if Path(save_to).exists():
        vectorstore.save_local(folder_path=save_to)

    print("> DONE")
    print(f"> Saved to: {save_to}")

    ### Save as zip, then remove faiss/pkl files to allow for upload to huggingface

    import shutil

    shutil.make_archive(save_to, 'zip', save_to)

    os.remove(save_to + "/index.faiss")
    os.remove(save_to + "/index.pkl")

    shutil.move(save_to + '.zip', save_to + "/" + save_to + '.zip')

    return vectorstore


# +
# https://colab.research.google.com/drive/1RWqGXd2B6sPchlYVihKaBSsHy9zWRcYF#scrollTo=Q_eTIZwf4Dk2

def docs_to_chroma_save(embeddings, docs_out:PandasDataFrame, save_to:str):
    print(f"> Total split documents: {len(docs_out)}")
    
    vectordb = Chroma.from_documents(documents=docs_out, 
                                 embedding=embeddings,
                                 persist_directory=save_to)
    
    # persiste the db to disk
    vectordb.persist()
    
    print("> DONE")
    print(f"> Saved to: {save_to}")
    
    return vectordb


# + [markdown] jp-MarkdownHeadingCollapsed=true
# ## Similarity search on saved vectorstore
# -

def sim_search_local_saved_vec(query, k_val, save_to="faiss_lambeth_census_embedding"):

    load_embeddings()

    docsearch = FAISS.load_local(folder_path=save_to, embeddings=embeddings)


    display(Markdown(question))

    search = docsearch.similarity_search_with_score(query, k=k_val)

    for item in search:
        print(item[0].page_content)
        print(f"Page: {item[0].metadata['source']}")
        print(f"Date: {item[0].metadata['date']}")
        print(f"Score: {item[1]}")
        print("---")
