# Import package

import os
from pathlib import Path
import re
import requests
import pandas as pd
import dateutil.parser
from typing import Type, List

from langchain_community.embeddings import HuggingFaceEmbeddings # HuggingFaceInstructEmbeddings, 
from langchain_community.vectorstores.faiss import FAISS
#from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document

from bs4 import BeautifulSoup
from docx import Document as Doc
from pypdf import PdfReader

PandasDataFrame = Type[pd.DataFrame]

split_strat = ["\n\n", "\n", ". ", "! ", "? "]
chunk_size = 300
chunk_overlap = 0
start_index = True

## Parse files
def determine_file_type(file_path):
        """
        Determine the file type based on its extension.
    
        Parameters:
            file_path (str): Path to the file.
    
        Returns:
            str: File extension (e.g., '.pdf', '.docx', '.txt', '.html').
        """
        return os.path.splitext(file_path)[1].lower()

def parse_file(file_paths, text_column='text'):
    """
    Accepts a list of file paths, determines each file's type based on its extension,
    and passes it to the relevant parsing function.
    
    Parameters:
        file_paths (list): List of file paths.
        text_column (str): Name of the column in CSV/Excel files that contains the text content.
    
    Returns:
        dict: A dictionary with file paths as keys and their parsed content (or error message) as values.
    """
    
    

    if not isinstance(file_paths, list):
        raise ValueError("Expected a list of file paths.")
    
    extension_to_parser = {
        '.pdf': parse_pdf,
        '.docx': parse_docx,
        '.txt': parse_txt,
        '.html': parse_html,
        '.htm': parse_html,  # Considering both .html and .htm for HTML files
        '.csv': lambda file_path: parse_csv_or_excel(file_path, text_column),
        '.xlsx': lambda file_path: parse_csv_or_excel(file_path, text_column)
    }
    
    parsed_contents = {}
    file_names = []

    for file_path in file_paths:
        print(file_path.name)
        #file = open(file_path.name, 'r')
        #print(file)
        file_extension = determine_file_type(file_path.name)
        if file_extension in extension_to_parser:
            parsed_contents[file_path.name] = extension_to_parser[file_extension](file_path.name)
        else:
            parsed_contents[file_path.name] = f"Unsupported file type: {file_extension}"

        filename_end = get_file_path_end(file_path.name)

        file_names.append(filename_end)
    
    return parsed_contents, file_names

def text_regex_clean(text):
    # Merge hyphenated words
        text = re.sub(r"(\w+)-\n(\w+)", r"\1\2", text)
        # If a double newline ends in a letter, add a full stop.
        text = re.sub(r'(?<=[a-zA-Z])\n\n', '.\n\n', text)
        # Fix newlines in the middle of sentences
        text = re.sub(r"(?<!\n\s)\n(?!\s\n)", " ", text.strip())
        # Remove multiple newlines
        text = re.sub(r"\n\s*\n", "\n\n", text)
        text = re.sub(r"  ", " ", text)
        # Add full stops and new lines between words with no space between where the second one has a capital letter
        text = re.sub(r'(?<=[a-z])(?=[A-Z])', '. \n\n', text)

        return text

def parse_csv_or_excel(file_paths, text_column = "text"):
        """
        Read in a CSV or Excel file.
        
        Parameters:
            file_path (str): Path to the CSV file.
            text_column (str): Name of the column in the CSV file that contains the text content.
        
        Returns:
            Pandas DataFrame: Dataframe output from file read
        """

        file_names = []
        out_df = pd.DataFrame()

        for file_path in file_paths:
            file_extension = determine_file_type(file_path.name)
            file_name = get_file_path_end(file_path.name)

            if file_extension == ".csv":
                df = pd.read_csv(file_path.name)
                if text_column not in df.columns: return pd.DataFrame(), ['Please choose a valid column name']
                df['source'] = file_name
                df['page_section'] = ""
            elif file_extension == ".xlsx":
                df = pd.read_excel(file_path.name, engine='openpyxl')
                if text_column not in df.columns: return pd.DataFrame(), ['Please choose a valid column name']
                df['source'] = file_name
                df['page_section'] = ""
            else:
                print(f"Unsupported file type: {file_extension}")
                return pd.DataFrame(), ['Please choose a valid file type']
            
            file_names.append(file_name)
            out_df = pd.concat([out_df, df])
        
        #if text_column not in df.columns:
        #    return f"Column '{text_column}' not found in {file_path}"
        #text_out = " ".join(df[text_column].dropna().astype(str))
        return out_df, file_names

def parse_excel(file_path, text_column):
        """
        Read text from an Excel file.
        
        Parameters:
            file_path (str): Path to the Excel file.
            text_column (str): Name of the column in the Excel file that contains the text content.
        
        Returns:
            Pandas DataFrame: Dataframe output from file read
        """
        df = pd.read_excel(file_path, engine='openpyxl')
        #if text_column not in df.columns:
        #    return f"Column '{text_column}' not found in {file_path}"
        #text_out = " ".join(df[text_column].dropna().astype(str))
        return df

def parse_pdf(file) -> List[str]:

    """
    Extract text from a PDF file.
    
    Parameters:
        file_path (str): Path to the PDF file.
    
    Returns:
        List[str]: Extracted text from the PDF.
    """
    
    output = []
    #for file in files:
    print(file) # .name
    pdf = PdfReader(file) #[i] .name[i]
    
    for page in pdf.pages:
            text = page.extract_text()
            
            text = text_regex_clean(text)

            output.append(text)
    return output

def parse_docx(file_path):
    """
    Reads the content of a .docx file and returns it as a string.

    Parameters:
    - file_path (str): Path to the .docx file.

    Returns:
    - str: Content of the .docx file.
    """
    doc = Doc(file_path)
    full_text = []
    for para in doc.paragraphs:
        para = text_regex_clean(para)

        full_text.append(para.text.replace("  ", " ").strip())
    return '\n'.join(full_text)

def parse_txt(file_path):
    """
    Read text from a TXT or HTML file.
    
    Parameters:
        file_path (str): Path to the TXT or HTML file.
    
    Returns:
        str: Text content of the file.
    """
    with open(file_path, 'r', encoding="utf-8") as file:
        file_contents = file.read().replace("  ", " ").strip()

        file_contents = text_regex_clean(file_contents)

        return file_contents

def parse_html(page_url, div_filter="p"):
    """
    Determine if the source is a web URL or a local HTML file, extract the content based on the div of choice. Also tries to extract dates (WIP)

    Parameters:
        page_url (str): The web URL or local file path.

    Returns:
        str: Extracted content.
    """

    def is_web_url(s):
        """
        Check if the input string is a web URL.
        """
        return s.startswith("http://") or s.startswith("https://")

    def is_local_html_file(s):
        """
        Check if the input string is a path to a local HTML file.
        """
        return (s.endswith(".html") or s.endswith(".htm")) and os.path.isfile(s)

    def extract_text_from_source(source):
        """
        Determine if the source is a web URL or a local HTML file, 
        and then extract its content accordingly.

        Parameters:
        source (str): The web URL or local file path.

        Returns:
        str: Extracted content.
        """
        if is_web_url(source):
            response = requests.get(source)
            response.raise_for_status()  # Raise an HTTPError for bad responses
            return response.text.replace("  ", " ").strip()
        elif is_local_html_file(source):
            with open(source, 'r', encoding='utf-8') as file:
                file_out = file.read().replace
                return file_out
        else:
            raise ValueError("Input is neither a valid web URL nor a local HTML file path.")
               

    def clean_html_data(data, date_filter="", div_filt="p"):
        """
        Extracts and cleans data from HTML content.

        Parameters:
            data (str): HTML content to be parsed.
            date_filter (str, optional): Date string to filter results. If set, only content with a date greater than this will be returned.
            div_filt (str, optional): HTML tag to search for text content. Defaults to "p".

        Returns:
            tuple: Contains extracted text and date as strings. Returns empty strings if not found.
        """
    
        soup = BeautifulSoup(data, 'html.parser')

        # Function to exclude div with id "bar"
        def exclude_div_with_id_bar(tag):
            return tag.has_attr('id') and tag['id'] == 'related-links'

        text_elements = soup.find_all(div_filt)
        date_elements = soup.find_all(div_filt, {"class": "page-neutral-intro__meta"})
    
        # Extract date
        date_out = ""
        if date_elements:
            date_out = re.search(">(.*?)<", str(date_elements[0])).group(1)
            date_dt = dateutil.parser.parse(date_out)

            if date_filter:
                date_filter_dt = dateutil.parser.parse(date_filter)
                if date_dt < date_filter_dt:
                    return '', date_out

        # Extract text
        text_out_final = ""
        if text_elements:
            text_out_final = '\n'.join(paragraph.text for paragraph in text_elements)
            text_out_final = text_regex_clean(text_out_final)
        else:
            print(f"No elements found with tag '{div_filt}'. No text returned.")
    
        return text_out_final, date_out
    

    #page_url = "https://pypi.org/project/InstructorEmbedding/" #'https://www.ons.gov.uk/visualisations/censusareachanges/E09000022/index.html'

    html_text = extract_text_from_source(page_url)
    #print(page.text)

    texts = []
    metadatas = []

    clean_text, date = clean_html_data(html_text, date_filter="", div_filt=div_filter)
    texts.append(clean_text)
    metadatas.append({"source": page_url, "date":str(date)})

    #print(metadatas)

    return texts, metadatas, page_url

def get_file_path_end(file_path):
    match = re.search(r'(.*[\/\\])?(.+)$', file_path)
        
    filename_end = match.group(2) if match else ''

    return filename_end

# +
# Convert parsed text to docs
# -

def text_to_docs(text_dict: dict, chunk_size: int = chunk_size) -> List[Document]:
    """
    Converts the output of parse_file (a dictionary of file paths to content)
    to a list of Documents with metadata.
    """
    
    doc_sections = []
    parent_doc_sections = []

    for file_path, content in text_dict.items():
        ext = os.path.splitext(file_path)[1].lower()

        # Depending on the file extension, handle the content
        if ext == '.pdf':
            docs, page_docs = pdf_text_to_docs(content, chunk_size)
        elif ext in ['.html', '.htm', '.txt', '.docx']:
            docs = html_text_to_docs(content, chunk_size)
        elif ext in ['.csv', '.xlsx']:
            docs, page_docs = csv_excel_text_to_docs(content, chunk_size)
        else:
            print(f"Unsupported file type {ext} for {file_path}. Skipping.")
            continue

        
        filename_end = get_file_path_end(file_path)

        #match = re.search(r'(.*[\/\\])?(.+)$', file_path)
        #filename_end = match.group(2) if match else ''

        # Add filename as metadata
        for doc in docs: doc.metadata["source"] = filename_end
        #for parent_doc in parent_docs: parent_doc.metadata["source"] = filename_end
        
        doc_sections.extend(docs)
        #parent_doc_sections.extend(parent_docs)

    return doc_sections#, page_docs

def pdf_text_to_docs(text, chunk_size: int = chunk_size) -> List[Document]:
    """Converts a string or list of strings to a list of Documents
    with metadata."""

    #print(text)

    if isinstance(text, str):
        # Take a single string as one page
        text = [text]
        
    page_docs = [Document(page_content=page, metadata={"page": page}) for page in text]
  

    # Add page numbers as metadata
    for i, doc in enumerate(page_docs):
        doc.metadata["page"] = i + 1

    print("page docs are: ")
    print(page_docs)

    # Split pages into sections
    doc_sections = []

    for doc in page_docs:

        #print("page content: ")
        #print(doc.page_content)

        if doc.page_content == '':
            sections = ['']

        else:
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                separators=split_strat,#["\n\n", "\n", ".", "!", "?", ",", " ", ""],
                chunk_overlap=chunk_overlap,
                add_start_index=True
            )
            sections = text_splitter.split_text(doc.page_content)
        
        for i, section in enumerate(sections):
            doc = Document(
                   page_content=section, metadata={"page": doc.metadata["page"], "section": i, "page_section": f"{doc.metadata['page']}-{i}"})

            
            doc_sections.append(doc)

    return doc_sections, page_docs#, parent_doc

def html_text_to_docs(texts, metadatas, chunk_size:int = chunk_size):

    text_splitter = RecursiveCharacterTextSplitter(
        separators=split_strat,#["\n\n", "\n", ".", "!", "?", ",", " ", ""],
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        add_start_index=True
    )

    #print(texts)
    #print(metadatas)

    documents = text_splitter.create_documents(texts, metadatas=metadatas)

    for i, section in enumerate(documents):
        section.metadata["page_section"] = i + 1

    

    return documents

def write_out_metadata_as_string(metadata_in):
    # If metadata_in is a single dictionary, wrap it in a list
    if isinstance(metadata_in, dict):
        metadata_in = [metadata_in]

    metadata_string = [f"{'  '.join(f'{k}: {v}' for k, v in d.items() if k != 'page_section')}" for d in metadata_in] # ['metadata']
    return metadata_string

def csv_excel_text_to_docs(df, text_column='text', chunk_size=None) -> List[Document]:
    """Converts a DataFrame's content to a list of Documents with metadata."""
    
    doc_sections = []
    df[text_column] = df[text_column].astype(str) # Ensure column is a string column

    # For each row in the dataframe
    for idx, row in df.iterrows():
        # Extract the text content for the document
        doc_content = row[text_column]
        
        # Generate metadata containing other columns' data
        metadata = {"row": idx + 1}
        for col, value in row.items():
            if col != text_column:
                metadata[col] = value

        metadata_string = write_out_metadata_as_string(metadata)[0]

        

        # If chunk_size is provided, split the text into chunks
        if chunk_size:
            # Assuming you have a text splitter function similar to the PDF handling
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                # Other arguments as required by the splitter
            )
            sections = text_splitter.split_text(doc_content)

            
            # For each section, create a Document object
            for i, section in enumerate(sections):
                section = '. '.join([metadata_string, section])
                doc = Document(page_content=section, 
                               metadata={**metadata, "section": i, "row_section": f"{metadata['row']}-{i}"})
                doc_sections.append(doc)
        else:
            # If no chunk_size is provided, create a single Document object for the row
            doc_content = '. '.join([metadata_string, doc_content])
            doc = Document(page_content=doc_content, metadata=metadata)
            doc_sections.append(doc)
    
    return doc_sections

# # Functions for working with documents after loading them back in

def pull_out_data(series):

    # define a lambda function to convert each string into a tuple
    to_tuple = lambda x: eval(x)

    # apply the lambda function to each element of the series
    series_tup = series.apply(to_tuple)

    series_tup_content = list(zip(*series_tup))[1]

    series = pd.Series(list(series_tup_content))#.str.replace("^Main post content", "", regex=True).str.strip()

    return series

def docs_from_csv(df):

    import ast
    
    documents = []
    
    page_content = pull_out_data(df["0"])
    metadatas = pull_out_data(df["1"])

    for x in range(0,len(df)):       
        new_doc = Document(page_content=page_content[x], metadata=metadatas[x])
        documents.append(new_doc)
        
    return documents

def docs_from_lists(docs, metadatas):

    documents = []

    for x, doc in enumerate(docs):
        new_doc = Document(page_content=doc, metadata=metadatas[x])
        documents.append(new_doc)
        
    return documents

def docs_elements_from_csv_save(docs_path="documents.csv"):

    documents = pd.read_csv(docs_path)

    docs_out = docs_from_csv(documents)

    out_df = pd.DataFrame(docs_out)

    docs_content = pull_out_data(out_df[0].astype(str))

    docs_meta = pull_out_data(out_df[1].astype(str))

    doc_sources = [d['source'] for d in docs_meta]

    return out_df, docs_content, docs_meta, doc_sources

# ## Create embeddings and save faiss vector store to the path specified in `save_to`

def load_embeddings(model_name = "BAAI/bge-base-en-v1.5"):

    #if model_name == "hkunlp/instructor-large":
    #    embeddings_func = HuggingFaceInstructEmbeddings(model_name=model_name,
    #    embed_instruction="Represent the paragraph for retrieval: ",
    #    query_instruction="Represent the question for retrieving supporting documents: "
    #    )

    #else: 
    embeddings_func = HuggingFaceEmbeddings(model_name=model_name)

    global embeddings

    embeddings = embeddings_func

    return embeddings_func

def embed_faiss_save_to_zip(docs_out, save_to="output", model_name = "BAAI/bge-base-en-v1.5"):

    load_embeddings(model_name=model_name)

    #embeddings_fast = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

    print(f"> Total split documents: {len(docs_out)}")

    vectorstore = FAISS.from_documents(documents=docs_out, embedding=embeddings)

    if not Path(save_to).exists(): 
        os.mkdir(save_to)

    if Path(save_to).exists():
        vectorstore.save_local(folder_path=save_to)

    print("> DONE")
    print(f"> Saved to: {save_to}")

    ### Save as zip, then remove faiss/pkl files to allow for upload to huggingface

    import shutil

    shutil.make_archive(save_to, 'zip', save_to)

    os.remove(save_to + "/index.faiss")
    os.remove(save_to + "/index.pkl")

    save_zip_out = save_to + "/" + save_to + '.zip'

    shutil.move(save_to + '.zip', save_zip_out)

    out_message = "Document processing complete"

    return out_message, vectorstore, save_zip_out

def docs_to_chroma_save(embeddings, docs_out:PandasDataFrame, save_to:str):
    print(f"> Total split documents: {len(docs_out)}")
    
    vectordb = Chroma.from_documents(documents=docs_out, 
                                 embedding=embeddings,
                                 persist_directory=save_to)
    
    # persiste the db to disk
    vectordb.persist()
    
    print("> DONE")
    print(f"> Saved to: {save_to}")
    
    return vectordb

def sim_search_local_saved_vec(query, k_val, save_to="faiss_lambeth_census_embedding"):

    load_embeddings()

    docsearch = FAISS.load_local(folder_path=save_to, embeddings=embeddings)


    display(Markdown(question))

    search = docsearch.similarity_search_with_score(query, k=k_val)

    for item in search:
        print(item[0].page_content)
        print(f"Page: {item[0].metadata['source']}")
        print(f"Date: {item[0].metadata['date']}")
        print(f"Score: {item[1]}")
        print("---")
